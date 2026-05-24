from __future__ import annotations

import io
import json
import logging
import re
import os
import base64
import subprocess
import sys
import time
import traceback
import threading
import urllib.error
import urllib.parse
import urllib.request
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from dotenv import load_dotenv
from pdf2image import convert_from_path, pdfinfo_from_path
from pdf2image.exceptions import PDFInfoNotInstalledError, PDFPageCountError, PDFSyntaxError
from PIL import Image
from utils import clear_blocking_proxy_env

try:
    from google import genai
    from google.genai import types
except Exception as exc:
    genai = None
    types = None
    GOOGLE_GENAI_IMPORT_ERROR = exc
else:
    GOOGLE_GENAI_IMPORT_ERROR = None


EXTRACTION_PROMPT = (
    "You are reading a scanned Ayurvedic medical textbook page. Return an exact visible transcription only. "
    "Treat the image as a collection of visual pixels, not as a language to be comprehended. Replicate the exact visible ink. "
    "Copy the text exactly as printed from the image, including Hindi, Sanskrit, English, headings, footnotes, tables, figure captions, page numbers, punctuation, spelling, capitalization, numbers, symbols, and medical terms. "
    "WARNING: The printed source text contains historical printing errors, misspelled Sanskrit, and grammatical typos (e.g., incorrect matras, wrong anusvaras, missing half-letters, like printing 'पुंष्यति' instead of 'पुष्यति' or 'लाहितादयः' instead of 'लोहितादयः'). YOU MUST PRESERVE THESE EXACT PRINTED ERRORS. "
    "Preserve the visible reading order and line/table structure as closely as plain Markdown text allows. "
    "For Sanskrit/Hindi shloka or quoted source verses, copy every Devanagari character exactly as visible, including matras, anusvara, visarga, half-letters, avagraha, punctuation, danda marks (। and ।।), reference text, and the same verse line breaks. "
    "Pay extremely close attention to the presence or absence of the avagraha symbol ('ऽ'). Only include it if it is physically printed on the page. "
    "Devanagari Matra Warning: Pay extreme attention to upper matras (e, ai, reph, anusvara). The ink may be blurry. Do not confuse the 'ऐ' (ai) matra for a 'र्' (reph). Do not insert an avagraha ('ऽ') unless there is a clear, definitive 'S' curve printed on the page. "
    "Do NOT apply autocorrect. Do NOT fix Sanskrit grammar or spelling. If the printed book has a typo, your transcription MUST have the exact same typo. "
    "Do not correct, modernize, complete, split, join, or reinterpret shloka text. "
    "For printed tables, preserve the same columns, rows, row order, cell text, blank cells, and stacked cell lines; do not move text between columns. "
    "Do not normalize spelling. Do not fix grammar. Do not expand abbreviations. Do not translate. Do not romanize. "
    "Do not join split words, remove printed hyphens, add missing words, repeat merged-cell text, add labels, add continuation notes, or clean formatting unless that exact text is visibly printed. "
    "If a character or word is unreadable from the image, write [unclear] in that exact position. "
    "Do not guess from context or medical knowledge. Do not summarize. Return only the transcription."
)

VERIFICATION_PROMPT = (
    "You are a strict medical textbook transcription verifier. Compare the extracted text against the scanned page image and return an exact visible transcription only. "
    "Treat the image as a collection of visual pixels, not as a language to be comprehended. Replicate the exact visible ink. "
    "Correct only characters, words, punctuation, line/table placement, and omissions that are directly visible in the image; never correct the source text itself. "
    "Preserve Hindi, Sanskrit, English, headings, tables, footnotes, figure captions, page numbers, spelling, capitalization, punctuation, numbers, symbols, and page order exactly as printed. "
    "Pay extremely close attention to the presence or absence of the avagraha symbol ('ऽ'). Only include it if it is physically printed on the page. "
    "Devanagari Matra Warning: Pay extreme attention to upper matras (e, ai, reph, anusvara). The ink may be blurry. Do not confuse the 'ऐ' (ai) matra for a 'र्' (reph). Do not insert an avagraha ('ऽ') unless there is a clear, definitive 'S' curve printed on the page. "
    "Do NOT apply autocorrect. Do NOT fix Sanskrit grammar or spelling. If the printed book has a typo, your transcription MUST have the exact same typo. "
    "WARNING: The printed source text contains historical printing errors, misspelled Sanskrit, and grammatical typos (e.g., incorrect matras, wrong anusvaras, missing half-letters, like printing 'पुंष्यति' instead of 'पुष्यति' or 'लाहितादयः' instead of 'लोहितादयः'). YOU MUST PRESERVE THESE EXACT PRINTED ERRORS. "
    "For Sanskrit/Hindi shloka or quoted source verses, copy every Devanagari character exactly as visible, including matras, anusvara, visarga, half-letters, avagraha, punctuation, danda marks (। and ।।), reference text, and the same verse line breaks. "
    "Do not correct, modernize, complete, split, join, or reinterpret shloka text. "
    "For printed tables, preserve the same columns, rows, row order, cell text, blank cells, and stacked cell lines; do not move text between columns. "
    "Do not normalize spelling. Do not fix grammar. Do not expand abbreviations. Do not translate. Do not romanize. "
    "Do not join split words, remove printed hyphens, add missing words, repeat merged-cell text, add labels, add continuation notes, or clean formatting unless that exact text is visibly printed. "
    "If a character or word is unreadable from the image, write [unclear] in that exact position. "
    "Do not guess from context or medical knowledge. Do not summarize. Return only the verified transcription."
)

FORMAT_PROMPT = """You are converting verified Ayurvedic textbook transcription into structured study notes. Use only the verified source text. Do not invent anything. Do not add modern medical explanation unless it is present in the source. Preserve Ayurvedic terms. If a section is not clearly present, write: Not clearly mentioned in source text.
Preserve the source languages exactly: Hindi should remain Hindi, Sanskrit should remain Sanskrit, and English should remain English. Do not translate Hindi or Sanskrit into English. Do not romanize Devanagari unless the verified source text itself is romanized.

Required Markdown format:

# Chapter: <chapter title>

## Source PDF
<filename>

## Verified Source Pages
Pages: <start page> to <end page>

## Nidan Panchak Notes

### 1. Nidana - Causes / Etiology
- ...

### 2. Purvarupa - Premonitory Symptoms
- ...

### 3. Rupa - Signs and Symptoms
- ...

### 4. Upashaya / Anupashaya - Relieving and Aggravating Factors
- ...

### 5. Samprapti - Pathogenesis
- ...

## Disease-wise Notes

For every disease/disorder mentioned in the source, create:

### <Disease / Disorder Name>

#### Definition / Description
- ...

#### Nidana
- ...

#### Purvarupa
- ...

#### Rupa
- ...

#### Upashaya / Anupashaya
- ...

#### Samprapti
- ...

#### Commentary Notes
- ...

#### Important Terms
- ...

#### Unclear / Needs Review
- ...

## Important Sanskrit / Ayurvedic Terms

| Term | Meaning | Context |
|------|---------|---------|

## Tables Found in Source
Recreate important source tables in Markdown only if present in the verified text.

## Figures / Diagrams Found in Source
Mention:
- figure number
- page number
- caption
- what the figure shows

## Key Clinical Takeaways
- ...

## Unclear / Needs Human Review
- ...
"""


STATE_IO_LOCK = threading.Lock()


class FatalGeminiError(RuntimeError):
    """A Gemini API/account error that cannot be fixed by retrying a page."""


@dataclass(frozen=True)
class Config:
    gemini_api_key: str
    gemini_model: str
    gemini_extraction_model: str
    gemini_verification_model: str
    gemini_format_model: str
    pdf_dir: Path
    output_dir: Path
    poppler_path: str | None
    dpi: int
    max_retries: int
    gemini_delay_seconds: float
    gemini_timeout_seconds: int
    speed_mode: str
    page_workers_per_job: int
    reset_master: bool
    create_docx: bool
    create_structured_notes: bool
    force_reprocess_pages: bool
    use_embedded_pdf_text: bool
    exact_text_only: bool
    test_start_page: int
    test_max_pages: int
    test_pdf_limit: int
    root_dir: Path
    extracted_dir: Path
    verified_dir: Path
    chapter_sources_dir: Path
    logs_dir: Path
    state_file: Path
    master_notes_file: Path


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def parse_bool(value: str | None, default: bool = False) -> bool:
    if value is None:
        return default
    return value.strip().lower() in {"1", "true", "yes", "y", "on"}


def parse_int(value: str | None, default: int, minimum: int = 1) -> int:
    try:
        parsed = int(str(value).strip()) if value is not None else default
    except ValueError:
        return default
    return max(parsed, minimum)


def parse_float(value: str | None, default: float, minimum: float = 0.0) -> float:
    try:
        parsed = float(str(value).strip()) if value is not None else default
    except ValueError:
        return default
    return max(parsed, minimum)


def load_config() -> Config:
    load_dotenv(override=True)
    clear_blocking_proxy_env()
    root_dir = Path.cwd()
    pdf_dir = Path(os.getenv("PDF_DIR", "./pdfs")).expanduser()
    output_dir = Path(os.getenv("OUTPUT_DIR", "./output_notes")).expanduser()

    if not pdf_dir.is_absolute():
        pdf_dir = root_dir / pdf_dir
    if not output_dir.is_absolute():
        output_dir = root_dir / output_dir

    poppler_path = os.getenv("POPPLER_PATH")
    if poppler_path:
        poppler_path = poppler_path.strip().strip('"') or None

    speed_mode = (os.getenv("SPEED_MODE", "accuracy").strip().lower() or "accuracy")
    if speed_mode not in {"fast", "balanced", "accuracy"}:
        speed_mode = "accuracy"

    return Config(
        gemini_api_key=os.getenv("GEMINI_API_KEY", "").strip(),
        gemini_model=os.getenv("GEMINI_MODEL", "gemini-2.5-flash").strip() or "gemini-2.5-flash",
        gemini_extraction_model=(
            os.getenv("GEMINI_EXTRACTION_MODEL")
            or os.getenv("GEMINI_MODEL")
            or "gemini-2.5-flash"
        ).strip(),
        gemini_verification_model=(
            os.getenv("GEMINI_VERIFICATION_MODEL")
            or os.getenv("GEMINI_MODEL")
            or "gemini-2.5-flash"
        ).strip(),
        gemini_format_model=(
            os.getenv("GEMINI_FORMAT_MODEL")
            or os.getenv("GEMINI_MODEL")
            or "gemini-2.5-flash"
        ).strip(),
        pdf_dir=pdf_dir,
        output_dir=output_dir,
        poppler_path=poppler_path,
        dpi=parse_int(os.getenv("DPI"), 250, minimum=72),
        max_retries=parse_int(os.getenv("MAX_RETRIES"), 5, minimum=1),
        gemini_delay_seconds=parse_float(os.getenv("GEMINI_DELAY_SECONDS"), 5.0, minimum=0.0),
        gemini_timeout_seconds=parse_int(os.getenv("GEMINI_TIMEOUT_SECONDS"), 180, minimum=30),
        speed_mode=speed_mode,
        page_workers_per_job=parse_int(os.getenv("PAGE_WORKERS_PER_JOB"), 1, minimum=1),
        reset_master=parse_bool(os.getenv("RESET_MASTER"), default=False),
        create_docx=parse_bool(os.getenv("CREATE_DOCX"), default=False),
        create_structured_notes=parse_bool(os.getenv("CREATE_STRUCTURED_NOTES"), default=False),
        force_reprocess_pages=parse_bool(os.getenv("FORCE_REPROCESS_PAGES"), default=False),
        use_embedded_pdf_text=parse_bool(os.getenv("USE_EMBEDDED_PDF_TEXT"), default=True),
        exact_text_only=parse_bool(os.getenv("EXACT_TEXT_ONLY"), default=False),
        test_start_page=parse_int(os.getenv("TEST_START_PAGE"), 1, minimum=1),
        test_max_pages=parse_int(os.getenv("TEST_MAX_PAGES"), 0, minimum=0),
        test_pdf_limit=parse_int(os.getenv("TEST_PDF_LIMIT"), 0, minimum=0),
        root_dir=root_dir,
        extracted_dir=root_dir / "extracted_pages",
        verified_dir=root_dir / "verified_pages",
        chapter_sources_dir=root_dir / "chapter_sources",
        logs_dir=root_dir / "logs",
        state_file=root_dir / "processing_state.json",
        master_notes_file=root_dir / "Complete_Book_Notes.md",
    )


def setup_logging(config: Config) -> None:
    config.logs_dir.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(config.logs_dir / "processor.log", encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )


def default_state() -> dict[str, Any]:
    return {
        "created_at": utc_now(),
        "updated_at": utc_now(),
        "completed_pdfs": [],
        "completed_pages": {},
        "extracted_pages": {},
        "verified_pages": {},
        "failed_pages": {},
        "formatted_chapters": [],
        "retry_counts": {},
        "chapter_page_counts": {},
        "master_appended_chapters": [],
    }


def normalize_state(state: dict[str, Any]) -> dict[str, Any]:
    base = default_state()
    for key, value in state.items():
        base[key] = value
    for key in [
        "completed_pdfs",
        "formatted_chapters",
        "master_appended_chapters",
    ]:
        if not isinstance(base.get(key), list):
            base[key] = []
    for key in [
        "completed_pages",
        "extracted_pages",
        "verified_pages",
        "failed_pages",
        "retry_counts",
        "chapter_page_counts",
    ]:
        if not isinstance(base.get(key), dict):
            base[key] = {}
    return base


def load_state(config: Config) -> dict[str, Any]:
    if not config.state_file.exists():
        return default_state()
    try:
        with config.state_file.open("r", encoding="utf-8") as handle:
            return normalize_state(json.load(handle))
    except json.JSONDecodeError:
        backup = config.state_file.with_suffix(f".invalid_{int(time.time())}.json")
        config.state_file.replace(backup)
        logging.error("Invalid processing_state.json moved to %s. Starting with a fresh state.", backup)
        return default_state()
    except Exception:
        logging.error("Could not load processing_state.json:\n%s", traceback.format_exc())
        return default_state()


def save_state(config: Config, state: dict[str, Any]) -> None:
    with STATE_IO_LOCK:
        state["updated_at"] = utc_now()
        temp_file = config.state_file.with_suffix(".json.tmp")
        with temp_file.open("w", encoding="utf-8") as handle:
            json.dump(state, handle, ensure_ascii=False, indent=2)
        temp_file.replace(config.state_file)


def ensure_output_dirs(config: Config) -> None:
    for directory in [
        config.pdf_dir,
        config.output_dir,
        config.extracted_dir,
        config.verified_dir,
        config.chapter_sources_dir,
        config.logs_dir,
    ]:
        directory.mkdir(parents=True, exist_ok=True)


def find_pdfs(config: Config) -> list[Path]:
    if not config.pdf_dir.exists():
        config.pdf_dir.mkdir(parents=True, exist_ok=True)
        print(f"Created missing PDF folder: {config.pdf_dir}")
        print("Put scanned textbook PDF files in this folder, then run the script again.")
        logging.warning("PDF folder was missing and has been created: %s", config.pdf_dir)
        return []
    pdfs = sorted([path for path in config.pdf_dir.iterdir() if path.is_file() and path.suffix.lower() == ".pdf"])
    skipped = sorted([path.name for path in config.pdf_dir.iterdir() if path.is_file() and path.suffix.lower() != ".pdf"])
    for name in skipped:
        logging.info("Skipping non-PDF file in pdfs folder: %s", name)
    return pdfs


def pdf2image_kwargs(config: Config) -> dict[str, Any]:
    kwargs: dict[str, Any] = {"dpi": config.dpi}
    if config.poppler_path:
        kwargs["poppler_path"] = config.poppler_path
    return kwargs


def get_pdf_page_count(pdf_path: Path, config: Config) -> int:
    kwargs = pdf2image_kwargs(config)
    kwargs.pop("dpi", None)
    try:
        info = pdfinfo_from_path(str(pdf_path), **kwargs)
        return int(info.get("Pages", 0))
    except PDFInfoNotInstalledError as exc:
        message = (
            "Poppler is not installed or POPPLER_PATH is wrong. On Windows, install Poppler and set "
            "POPPLER_PATH=C:\\poppler\\bin in .env, then run again."
        )
        raise RuntimeError(message) from exc
    except (PDFPageCountError, PDFSyntaxError) as exc:
        raise RuntimeError(f"Could not read PDF pages. The file may be broken or encrypted: {pdf_path.name}") from exc
    except Exception as exc:
        raise RuntimeError(f"Could not inspect PDF page count for {pdf_path.name}: {exc}") from exc


def convert_pdf_page_to_image(pdf_path: Path, config: Config, page_number: int) -> Image.Image:
    kwargs = pdf2image_kwargs(config)
    kwargs["first_page"] = page_number
    kwargs["last_page"] = page_number
    try:
        images = convert_from_path(str(pdf_path), **kwargs)
        if not images:
            raise RuntimeError(f"No image was produced for page {page_number}")
        return images[0]
    except PDFInfoNotInstalledError as exc:
        message = (
            "Poppler is not installed or POPPLER_PATH is wrong. On Windows, install Poppler and set "
            "POPPLER_PATH=C:\\poppler\\bin in .env, then run again."
        )
        raise RuntimeError(message) from exc
    except (PDFPageCountError, PDFSyntaxError) as exc:
        raise RuntimeError(f"Could not read PDF pages. The file may be broken or encrypted: {pdf_path.name}") from exc
    except Exception as exc:
        raise RuntimeError(f"PDF page conversion failed for {pdf_path.name} page {page_number}: {exc}") from exc


def pdftotext_exe(config: Config) -> str:
    if config.poppler_path:
        candidate = Path(config.poppler_path) / "pdftotext.exe"
        if candidate.exists():
            return str(candidate)
    return "pdftotext"


def extract_embedded_pdf_page_text(pdf_path: Path, config: Config, page_number: int) -> str:
    temp_dir = config.root_dir / ".tmp"
    temp_dir.mkdir(parents=True, exist_ok=True)
    output_path = temp_dir / f"{pdf_path.stem}_page_{page_number:03d}_pdftotext.txt"
    command = [
        pdftotext_exe(config),
        "-f",
        str(page_number),
        "-l",
        str(page_number),
        "-layout",
        "-enc",
        "UTF-8",
        str(pdf_path),
        str(output_path),
    ]
    try:
        completed = subprocess.run(command, capture_output=True, text=True, timeout=60, check=False)
    except FileNotFoundError:
        logging.warning("pdftotext is not available, falling back to OCR for %s page %03d.", pdf_path.name, page_number)
        return ""
    except subprocess.TimeoutExpired:
        logging.warning("pdftotext timed out for %s page %03d.", pdf_path.name, page_number)
        return ""

    if completed.returncode != 0:
        logging.warning(
            "pdftotext failed for %s page %03d: %s",
            pdf_path.name,
            page_number,
            completed.stderr.strip(),
        )
        return ""
    if not output_path.exists():
        return ""
    text = output_path.read_text(encoding="utf-8", errors="replace")
    return text.replace("\f", "").rstrip()


def pil_image_to_bytes(image: Image.Image) -> bytes:
    buffer = io.BytesIO()
    if image.mode not in {"RGB", "L"}:
        image = image.convert("RGB")
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def is_retryable_error(error: Exception) -> bool:
    text = str(error).lower()
    retry_terms = [
        "429",
        "quota",
        "rate limit",
        "resource exhausted",
        "too many requests",
        "503",
        "unavailable",
        "deadline",
        "timeout",
        "temporarily",
        "connection",
        "network",
    ]
    return any(term in text for term in retry_terms)


def is_fatal_gemini_error(error: Exception) -> bool:
    text = str(error).lower()
    fatal_terms = [
        "401",
        "403",
        "forbidden",
        "permission_denied",
        "permission denied",
        "api key not valid",
        "invalid api key",
        "dunning decision",
        "billing",
        "suspended",
    ]
    return any(term in text for term in fatal_terms)


def response_text(response: Any) -> str:
    text = getattr(response, "text", None)
    if text:
        return str(text).strip()
    try:
        candidates = getattr(response, "candidates", None) or []
        parts: list[str] = []
        for candidate in candidates:
            content = getattr(candidate, "content", None)
            for part in getattr(content, "parts", []) or []:
                part_text = getattr(part, "text", None)
                if part_text:
                    parts.append(str(part_text))
        return "\n".join(parts).strip()
    except Exception:
        return ""


def rest_response_text(payload: dict[str, Any]) -> str:
    parts: list[str] = []
    for candidate in payload.get("candidates", []) or []:
        content = candidate.get("content", {}) if isinstance(candidate, dict) else {}
        for part in content.get("parts", []) or []:
            if isinstance(part, dict) and part.get("text"):
                parts.append(str(part["text"]))
    return "\n".join(parts).strip()


def rest_part(content: Any) -> dict[str, Any]:
    if isinstance(content, str):
        return {"text": content}
    if isinstance(content, dict):
        if "inline_data" in content or "text" in content:
            return content
        if "inlineData" in content:
            return {"inline_data": content["inlineData"]}
    text = getattr(content, "text", None)
    if text:
        return {"text": str(text)}
    inline_data = getattr(content, "inline_data", None) or getattr(content, "inlineData", None)
    if inline_data:
        data = getattr(inline_data, "data", None)
        mime_type = getattr(inline_data, "mime_type", None) or getattr(inline_data, "mimeType", None)
        if data and mime_type:
            if isinstance(data, bytes):
                data = base64.b64encode(data).decode("ascii")
            return {"inline_data": {"mime_type": str(mime_type), "data": str(data)}}
    raise TypeError(f"Unsupported Gemini content part: {type(content).__name__}")


class RestGeminiResponse:
    def __init__(self, payload: dict[str, Any]) -> None:
        self.payload = payload
        self.text = rest_response_text(payload)


class RestGeminiModels:
    def __init__(self, client: "RestGeminiClient") -> None:
        self.client = client

    def generate_content(self, model: str, contents: Any, config: Any = None) -> RestGeminiResponse:
        return self.client.generate_content(model=model, contents=contents, config=config)


class RestGeminiClient:
    def __init__(self, api_key: str, timeout_seconds: int) -> None:
        self.api_key = api_key
        self.timeout_seconds = timeout_seconds
        self.models = RestGeminiModels(self)
        self.opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))

    def generate_content(self, model: str, contents: Any, config: Any = None) -> RestGeminiResponse:
        items = contents if isinstance(contents, list) else [contents]
        parts = [rest_part(item) for item in items]
        payload: dict[str, Any] = {
            "contents": [{"role": "user", "parts": parts}],
            "generationConfig": {"temperature": 0, "topP": 0.1},
        }
        if isinstance(config, dict):
            payload["generationConfig"].update(config)

        model_path = urllib.parse.quote(model, safe="")
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_path}:generateContent?key={self.api_key}"
        request = urllib.request.Request(
            url,
            data=json.dumps(payload).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with self.opener.open(request, timeout=self.timeout_seconds) as response:
                response_payload = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            body = exc.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"Gemini REST HTTP {exc.code}: {body}") from exc
        if isinstance(response_payload, dict) and response_payload.get("error"):
            raise RuntimeError(f"Gemini REST error: {response_payload['error']}")
        return RestGeminiResponse(response_payload)


def gemini_client(config: Config) -> genai.Client:
    if genai is not None and types is not None:
        return genai.Client(
            api_key=config.gemini_api_key,
            http_options=types.HttpOptions(timeout=config.gemini_timeout_seconds * 1000),
        )
    logging.warning("google-genai unavailable; using Gemini REST fallback: %s", GOOGLE_GENAI_IMPORT_ERROR)
    return RestGeminiClient(config.gemini_api_key, config.gemini_timeout_seconds)


def call_gemini_with_retry(
    client: genai.Client,
    config: Config,
    state: dict[str, Any],
    retry_key: str,
    contents: list[Any],
    model: str | None = None,
) -> str:
    last_error: Exception | None = None
    model_name = model or config.gemini_model
    for attempt in range(1, config.max_retries + 1):
        if config.gemini_delay_seconds:
            time.sleep(config.gemini_delay_seconds)
        try:
            generation_config: Any
            if types is not None:
                generation_config = types.GenerateContentConfig(
                    temperature=0,
                    top_p=0.1,
                )
            else:
                generation_config = {"temperature": 0, "topP": 0.1}
            response = client.models.generate_content(
                model=model_name,
                contents=contents,
                config=generation_config,
            )
            text = response_text(response)
            if not text:
                raise RuntimeError("Empty Gemini response")
            return text
        except Exception as exc:
            last_error = exc
            state.setdefault("retry_counts", {})
            state["retry_counts"][retry_key] = int(state["retry_counts"].get(retry_key, 0)) + 1
            save_state(config, state)
            if is_fatal_gemini_error(exc):
                message = (
                    f"Gemini API rejected {retry_key}: {exc}. "
                    "This is an API key, project permission, or billing problem, not a page OCR failure. "
                    "Fix GEMINI_API_KEY / Google project billing and permissions, then run the job again."
                )
                logging.error(message)
                raise FatalGeminiError(message) from exc
            retryable = is_retryable_error(exc)
            logging.warning(
                "Gemini call failed for %s on attempt %s/%s: %s",
                f"{retry_key} ({model_name})",
                attempt,
                config.max_retries,
                exc,
            )
            if attempt >= config.max_retries:
                break
            sleep_seconds = 60 if retryable else min(60, 2**attempt)
            if retryable:
                print(f"Gemini is rate-limited or unavailable for {retry_key}. Sleeping {sleep_seconds} seconds before retry.")
            time.sleep(sleep_seconds)
    raise RuntimeError(f"Gemini failed after {config.max_retries} attempts for {retry_key}: {last_error}")


def image_part(image_bytes: bytes) -> types.Part:
    if types is not None:
        return types.Part.from_bytes(data=image_bytes, mime_type="image/png")
    return {
        "inline_data": {
            "mime_type": "image/png",
            "data": base64.b64encode(image_bytes).decode("ascii"),
        }
    }


def extract_page_text(
    client: genai.Client,
    config: Config,
    state: dict[str, Any],
    pdf_stem: str,
    page_number: int,
    image: Image.Image,
) -> str:
    image_bytes = pil_image_to_bytes(image)
    retry_key = f"{pdf_stem}:page_{page_number:03d}:extract"
    text = call_gemini_with_retry(
        client,
        config,
        state,
        retry_key,
        [EXTRACTION_PROMPT, image_part(image_bytes)],
        model=config.gemini_extraction_model,
    )
    
    roman_chars = len(re.findall(r"[a-zA-Z]", text))
    devanagari_chars = len(re.findall(r"[\u0900-\u097F]", text))
    if roman_chars > devanagari_chars:
        raise ValueError("OCR Failure: More English characters than Devanagari found.")
        
    return text


def verify_page_text(
    client: genai.Client,
    config: Config,
    state: dict[str, Any],
    pdf_stem: str,
    page_number: int,
    image: Image.Image,
    extracted_text: str,
) -> str:
    image_bytes = pil_image_to_bytes(image)
    retry_key = f"{pdf_stem}:page_{page_number:03d}:verify"
    prompt = f"{VERIFICATION_PROMPT}\n\nExtracted text to verify:\n\n{extracted_text}"
    return call_gemini_with_retry(
        client,
        config,
        state,
        retry_key,
        [prompt, image_part(image_bytes)],
        model=config.gemini_verification_model,
    )


def page_file(directory: Path, page_number: int, suffix: str) -> Path:
    return directory / f"page_{page_number:03d}_{suffix}.md"


def int_list(state: dict[str, Any], section: str, pdf_stem: str) -> list[int]:
    values = state.setdefault(section, {}).setdefault(pdf_stem, [])
    cleaned: list[int] = []
    for value in values:
        try:
            cleaned.append(int(value))
        except (TypeError, ValueError):
            continue
    return sorted(set(cleaned))


def add_page_state(config: Config, state: dict[str, Any], section: str, pdf_stem: str, page_number: int) -> None:
    pages = set(int_list(state, section, pdf_stem))
    pages.add(page_number)
    state.setdefault(section, {})[pdf_stem] = sorted(pages)
    save_state(config, state)


def mark_page_extracted(config: Config, state: dict[str, Any], pdf_stem: str, page_number: int) -> None:
    add_page_state(config, state, "extracted_pages", pdf_stem, page_number)


def mark_page_verified(config: Config, state: dict[str, Any], pdf_stem: str, page_number: int) -> None:
    add_page_state(config, state, "verified_pages", pdf_stem, page_number)
    add_page_state(config, state, "completed_pages", pdf_stem, page_number)
    failed = state.setdefault("failed_pages", {}).setdefault(pdf_stem, {})
    failed.pop(str(page_number), None)
    if not failed:
        state.get("failed_pages", {}).pop(pdf_stem, None)
    save_state(config, state)


def mark_page_failed(
    config: Config,
    state: dict[str, Any],
    pdf_stem: str,
    page_number: int,
    stage: str,
    error: Exception,
) -> None:
    failed = state.setdefault("failed_pages", {}).setdefault(pdf_stem, {})
    failed[str(page_number)] = {
        "stage": stage,
        "error": str(error),
        "status": "NEEDS_REVIEW",
        "timestamp": utc_now(),
    }
    save_state(config, state)


def failed_page_reason(state: dict[str, Any], pdf_stem: str, page_number: int) -> str:
    failed = state.get("failed_pages", {}).get(pdf_stem, {})
    if not isinstance(failed, dict):
        return "Gemini could not verify this page automatically"
    info = failed.get(str(page_number), {})
    if not isinstance(info, dict):
        return "Gemini could not verify this page automatically"
    return str(info.get("error") or "Gemini could not verify this page automatically")


def page_review_placeholder(page_number: int, reason: str) -> str:
    return f"[PAGE {page_number} NEEDS MANUAL REVIEW: {reason}]"


def write_review_report(
    config: Config,
    state: dict[str, Any],
    pdf_stem: str,
    pdf_name: str,
    page_numbers: list[int],
) -> Path | None:
    failed = state.get("failed_pages", {}).get(pdf_stem, {})
    if not isinstance(failed, dict) or not failed:
        return None

    verified_pages = set(int_list(state, "verified_pages", pdf_stem))
    report_path = config.root_dir / "review_report.md"
    lines = [
        "# Review Report",
        "",
        f"PDF: {pdf_name}",
        "",
        "## Summary",
        f"- total pages: {len(page_numbers)}",
        f"- verified pages: {len([page for page in page_numbers if page in verified_pages])}",
        f"- failed pages: {len(failed)}",
        f"- failed page numbers: {', '.join(sorted(failed.keys(), key=lambda value: int(value) if value.isdigit() else 999999))}",
        "",
        "## Failed Pages",
        "",
    ]
    for page_text in sorted(failed.keys(), key=lambda value: int(value) if value.isdigit() else 999999):
        info = failed.get(page_text, {})
        row = info if isinstance(info, dict) else {}
        lines.extend(
            [
                f"### Page {page_text}",
                f"- stage: {row.get('stage', 'page_processing')}",
                f"- error: {str(row.get('error', '')).replace(chr(10), ' ')}",
                "",
            ]
        )
    report_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    logging.warning("Review report created for %s: %s", pdf_name, report_path)
    return report_path


def mark_pdf_complete(config: Config, state: dict[str, Any], pdf_stem: str) -> None:
    completed = set(state.setdefault("completed_pdfs", []))
    completed.add(pdf_stem)
    state["completed_pdfs"] = sorted(completed)
    save_state(config, state)


def combine_verified_pages(
    config: Config,
    pdf_stem: str,
    page_numbers: list[int],
    output_stem: str | None = None,
    allow_missing: bool = False,
) -> Path:
    source_name = output_stem or pdf_stem
    source_path = config.chapter_sources_dir / f"{source_name}_verified_source.md"
    verified_pdf_dir = config.verified_dir / pdf_stem
    state = load_state(config) if allow_missing else {}
    pieces: list[str] = []
    for page_number in page_numbers:
        verified_path = page_file(verified_pdf_dir, page_number, "verified")
        if not verified_path.exists():
            if allow_missing:
                reason = failed_page_reason(state, pdf_stem, page_number)
                pieces.append(
                    f"<!-- PAGE {page_number} START -->\n"
                    f"{page_review_placeholder(page_number, reason)}\n"
                    f"<!-- PAGE {page_number} END -->"
                )
                continue
            raise RuntimeError(f"Missing verified page file: {verified_path}")
        text = verified_path.read_text(encoding="utf-8").strip()
        pieces.append(f"<!-- PAGE {page_number} START -->\n{text}\n<!-- PAGE {page_number} END -->")
    source_path.write_text("\n\n".join(pieces).strip() + "\n", encoding="utf-8")
    return source_path


def format_chapter_notes(
    client: genai.Client,
    config: Config,
    state: dict[str, Any],
    pdf_path: Path,
    source_path: Path,
    page_range_text: str,
) -> str:
    source_text = source_path.read_text(encoding="utf-8")
    prompt = (
        f"{FORMAT_PROMPT}\n\n"
        f"Source PDF filename: {pdf_path.name}\n"
        f"Verified page range: {page_range_text}\n\n"
        f"Verified source text:\n\n{source_text}"
    )
    return call_gemini_with_retry(
        client,
        config,
        state,
        f"{pdf_path.stem}:format_chapter",
        [prompt],
        model=config.gemini_format_model,
    )


def chapter_already_in_master(config: Config, pdf_name: str) -> bool:
    if not config.master_notes_file.exists():
        return False
    marker = f"# END OF CHAPTER: {pdf_name}"
    try:
        return marker in config.master_notes_file.read_text(encoding="utf-8")
    except Exception:
        logging.warning("Could not inspect master notes file. State file will be used for append decision.")
        return False


def append_to_master_notes(config: Config, state: dict[str, Any], pdf_name: str, notes: str) -> None:
    appended = set(state.setdefault("master_appended_chapters", []))
    if pdf_name in appended or chapter_already_in_master(config, pdf_name):
        logging.info("Master notes already contain %s. Skipping append.", pdf_name)
        return

    with config.master_notes_file.open("a", encoding="utf-8") as handle:
        if config.master_notes_file.stat().st_size:
            handle.write("\n\n")
        handle.write(notes.strip())
        handle.write(f"\n\n---\n\n# END OF CHAPTER: {pdf_name}\n\n---\n")
    appended.add(pdf_name)
    state["master_appended_chapters"] = sorted(appended)
    save_state(config, state)


def create_docx_optional(
    config: Config,
    pdf_stem: str,
    page_numbers: list[int],
    output_stem: str | None = None,
    allow_missing: bool = False,
) -> Path | None:
    if not config.create_docx:
        return None
    try:
        from docx import Document
    except Exception as exc:
        logging.error("CREATE_DOCX=true, but python-docx is unavailable: %s", exc)
        return None

    docx_stem = output_stem or f"{pdf_stem}_verified"
    document = Document()
    document.add_heading(pdf_stem, level=1)
    verified_pdf_dir = config.verified_dir / pdf_stem
    state = load_state(config) if allow_missing else {}
    for page_number in page_numbers:
        verified_path = page_file(verified_pdf_dir, page_number, "verified")
        if not verified_path.exists():
            if allow_missing:
                document.add_heading(f"Page {page_number}", level=2)
                reason = failed_page_reason(state, pdf_stem, page_number)
                document.add_paragraph(page_review_placeholder(page_number, reason))
            continue
        document.add_heading(f"Page {page_number}", level=2)
        text = verified_path.read_text(encoding="utf-8").strip()
        add_markdown_text_to_docx(document, text)
    output_path = config.output_dir / f"{docx_stem}.docx"
    document.save(output_path)
    logging.info("DOCX created from verified text: %s", output_path)
    return output_path


def is_markdown_table_line(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def is_markdown_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    cells = [cell.strip() for cell in stripped.split("|")]
    return all(cell and set(cell) <= {"-", ":", " "} for cell in cells)


def markdown_table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_cell_text(cell: Any, text: str) -> None:
    paragraphs = cell.paragraphs
    paragraph = paragraphs[0] if paragraphs else cell.add_paragraph()
    parts = text.replace("<br/>", "<br>").replace("<br />", "<br>").split("<br>")
    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break()
        paragraph.add_run(part.strip())


def add_markdown_table_to_docx(document: Any, table_lines: list[str]) -> None:
    rows = [markdown_table_cells(line) for line in table_lines if not is_markdown_table_separator(line)]
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell_text = row[column_index] if column_index < len(row) else ""
            add_cell_text(table.cell(row_index, column_index), cell_text)
        if row_index == 0:
            for cell in table.rows[row_index].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_markdown_line_to_docx(document: Any, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    if stripped.startswith("# "):
        document.add_heading(stripped[2:].strip(), level=1)
    elif stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=2)
    elif stripped.startswith("### "):
        document.add_heading(stripped[4:].strip(), level=3)
    elif stripped.startswith("#### "):
        document.add_heading(stripped[5:].strip(), level=4)
    elif stripped.startswith("- "):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
    elif stripped == "---":
        document.add_paragraph("")
    else:
        document.add_paragraph(stripped)


def add_markdown_text_to_docx(document: Any, text: str) -> None:
    table_buffer: list[str] = []
    for line in text.splitlines():
        if is_markdown_table_line(line):
            table_buffer.append(line)
            continue
        if table_buffer:
            add_markdown_table_to_docx(document, table_buffer)
            table_buffer = []
        add_markdown_line_to_docx(document, line)
    if table_buffer:
        add_markdown_table_to_docx(document, table_buffer)


def create_formatted_notes_docx(config: Config, notes_path: Path) -> Path | None:
    if not config.create_docx:
        return None
    try:
        from docx import Document
    except Exception as exc:
        logging.error("CREATE_DOCX=true, but python-docx is unavailable: %s", exc)
        return None

    docx_path = notes_path.with_suffix(".docx")
    document = Document()
    add_markdown_text_to_docx(document, notes_path.read_text(encoding="utf-8"))
    document.save(docx_path)
    logging.info("DOCX created from formatted notes: %s", docx_path)
    return docx_path


def is_chapter_fully_verified(state: dict[str, Any], pdf_stem: str, page_numbers: list[int]) -> bool:
    verified = set(int_list(state, "verified_pages", pdf_stem))
    return all(page_number in verified for page_number in page_numbers)


def process_pdf(client: genai.Client, config: Config, state: dict[str, Any], pdf_path: Path) -> None:
    pdf_stem = pdf_path.stem
    recorded_failures = state.get("failed_pages", {}).get(pdf_stem, {})
    has_recorded_failures = isinstance(recorded_failures, dict) and bool(recorded_failures)
    if not config.test_max_pages and pdf_stem in set(state.get("completed_pdfs", [])) and not has_recorded_failures:
        logging.info("Skipping completed PDF before conversion: %s", pdf_path.name)
        print(f"Skipping completed PDF: {pdf_path.name}")
        return

    logging.info("Starting PDF: %s", pdf_path.name)
    print(f"\nProcessing PDF: {pdf_path.name}")

    try:
        total_pages = get_pdf_page_count(pdf_path, config)
    except Exception as exc:
        logging.error("%s", exc)
        state.setdefault("failed_pages", {}).setdefault(pdf_stem, {})["PDF"] = {
            "stage": "get_pdf_page_count",
            "error": str(exc),
            "timestamp": utc_now(),
        }
        save_state(config, state)
        return

    if config.test_max_pages:
        start_page = min(config.test_start_page, total_pages)
        end_page = min(total_pages, start_page + config.test_max_pages - 1)
        page_numbers = list(range(start_page, end_page + 1))
    else:
        page_numbers = list(range(1, total_pages + 1))
    pages_to_process = len(page_numbers)
    page_range_text = f"{page_numbers[0]} to {page_numbers[-1]}" if page_numbers else "none"
    state.setdefault("chapter_page_counts", {})[pdf_stem] = total_pages
    save_state(config, state)

    if total_pages == 0:
        logging.error("No pages were produced for PDF: %s", pdf_path.name)
        return
    if config.test_max_pages:
        logging.info(
            "Test mode active for %s: processing pages %s of %s total pages.",
            pdf_path.name,
            page_range_text,
            total_pages,
        )
        print(f"Test mode: processing pages {page_range_text} of {total_pages} total pages.")

    raw_pdf_dir = config.extracted_dir / pdf_stem
    verified_pdf_dir = config.verified_dir / pdf_stem
    raw_pdf_dir.mkdir(parents=True, exist_ok=True)
    verified_pdf_dir.mkdir(parents=True, exist_ok=True)
    state_lock = threading.Lock()
    fatal_error_event = threading.Event()

    def process_page(index: int) -> None:
        if fatal_error_event.is_set():
            return

        verified_path = page_file(verified_pdf_dir, index, "verified")
        raw_path = page_file(raw_pdf_dir, index, "raw")

        if (
            not config.force_reprocess_pages
            and verified_path.exists()
            and index in set(int_list(state, "verified_pages", pdf_stem))
        ):
            logging.info("Skipping already verified page: %s page %03d", pdf_stem, index)
            return

        logging.info("Starting page: %s page %03d", pdf_stem, index)
        print(f"Page {index}/{total_pages}")

        try:
            page_client = gemini_client(config)
            embedded_text = ""
            if config.use_embedded_pdf_text:
                embedded_text = extract_embedded_pdf_page_text(pdf_path, config, index)

            if embedded_text.strip():
                raw_path.write_text(embedded_text.strip() + "\n", encoding="utf-8")
                verified_path.write_text(embedded_text.strip() + "\n", encoding="utf-8")
                with state_lock:
                    mark_page_extracted(config, state, pdf_stem, index)
                    mark_page_verified(config, state, pdf_stem, index)
                logging.info("Used embedded PDF text without OCR: %s page %03d", pdf_stem, index)
                return

            if config.exact_text_only:
                raise RuntimeError(
                    "No embedded PDF text found on this page. Exact text-only mode will not create OCR text from scanned pixels."
                )

            image = convert_pdf_page_to_image(pdf_path, config, index)
            if raw_path.exists() and not config.force_reprocess_pages:
                extracted_text = raw_path.read_text(encoding="utf-8")
                with state_lock:
                    mark_page_extracted(config, state, pdf_stem, index)
                logging.info("Using existing raw extraction: %s page %03d", pdf_stem, index)
            else:
                extracted_text = extract_page_text(page_client, config, state, pdf_stem, index, image)
                raw_path.write_text(extracted_text.strip() + "\n", encoding="utf-8")
                with state_lock:
                    mark_page_extracted(config, state, pdf_stem, index)
                logging.info("Extraction complete: %s page %03d", pdf_stem, index)

            if config.speed_mode == "fast":
                verified_text = extracted_text
                logging.info("Fast mode: skipped second-pass verification for %s page %03d", pdf_stem, index)
            else:
                verified_text = verify_page_text(page_client, config, state, pdf_stem, index, image, extracted_text)
            verified_path.write_text(verified_text.strip() + "\n", encoding="utf-8")
            with state_lock:
                mark_page_verified(config, state, pdf_stem, index)
            logging.info("Verification complete: %s page %03d", pdf_stem, index)
        except KeyboardInterrupt:
            raise
        except FatalGeminiError:
            fatal_error_event.set()
            logging.error("Stopping %s because Gemini API is not usable:\n%s", pdf_stem, traceback.format_exc())
            raise
        except Exception as exc:
            logging.error("Failed %s page %03d:\n%s", pdf_stem, index, traceback.format_exc())
            with state_lock:
                mark_page_failed(config, state, pdf_stem, index, "page_processing", exc)
            print(f"Failed page {index}. Progress was saved; check logs/processor.log and processing_state.json.")

    workers = min(config.page_workers_per_job, len(page_numbers))
    try:
        if workers <= 1:
            for index in page_numbers:
                process_page(index)
        else:
            logging.info("Processing %s with %s page workers in %s mode.", pdf_path.name, workers, config.speed_mode)
            executor = ThreadPoolExecutor(max_workers=workers)
            futures = [executor.submit(process_page, index) for index in page_numbers]
            try:
                for future in as_completed(futures):
                    future.result()
            except FatalGeminiError:
                fatal_error_event.set()
                for future in futures:
                    future.cancel()
                executor.shutdown(wait=False, cancel_futures=True)
                raise
            else:
                executor.shutdown(wait=True)
    except FatalGeminiError as exc:
        state["api_error"] = {
            "pdf": pdf_path.name,
            "error": str(exc),
            "timestamp": utc_now(),
        }
        save_state(config, state)
        print("Gemini API rejected the request. Fix the API key/project billing, then click Run Again / Resume.")
        raise

    has_failed_pages = not is_chapter_fully_verified(state, pdf_stem, page_numbers)
    if has_failed_pages:
        write_review_report(config, state, pdf_stem, pdf_path.name, page_numbers)
        logging.warning("PDF has unverified failed pages. Creating DOCX with review placeholders: %s", pdf_path.name)
        print(f"Creating DOCX for {pdf_path.name}; failed pages will be marked for manual review.")

    try:
        output_stem = f"{pdf_stem}_TEST_pages_{page_numbers[0]}_to_{page_numbers[-1]}" if config.test_max_pages else pdf_stem
        source_path = combine_verified_pages(
            config,
            pdf_stem,
            page_numbers,
            output_stem=output_stem,
            allow_missing=has_failed_pages,
        )
        logging.info("Verified source created: %s", source_path)

        exact_output_stem = (
            f"{pdf_stem}_TEST_pages_{page_numbers[0]}_to_{page_numbers[-1]}_exact_verified_transcription"
            if config.test_max_pages
            else f"{pdf_stem}_exact_verified_transcription"
        )
        verified_docx_path = create_docx_optional(
            config,
            pdf_stem,
            page_numbers,
            output_stem=exact_output_stem,
            allow_missing=has_failed_pages,
        )

        notes_path: Path | None = None
        docx_notes_path: Path | None = None
        final_notes = source_path.read_text(encoding="utf-8")
        if config.create_structured_notes and has_failed_pages:
            logging.warning(
                "CREATE_STRUCTURED_NOTES=true, but structured notes were skipped because %s has failed pages.",
                pdf_path.name,
            )
            print("Structured notes skipped until all pages are verified.")
        elif config.create_structured_notes:
            formatted_chapters = set(state.setdefault("formatted_chapters", []))
            notes_path = config.output_dir / f"{output_stem}_nidan_panchak_notes.md"
            if config.test_max_pages and notes_path.exists():
                logging.info("Skipping existing test formatted notes: %s", notes_path)
            elif not config.test_max_pages and pdf_stem in formatted_chapters and notes_path.exists():
                logging.info("Skipping already formatted chapter: %s", pdf_stem)
            else:
                notes = format_chapter_notes(client, config, state, pdf_path, source_path, page_range_text)
                notes_path.write_text(notes.strip() + "\n", encoding="utf-8")
                if not config.test_max_pages:
                    formatted_chapters.add(pdf_stem)
                    state["formatted_chapters"] = sorted(formatted_chapters)
                    save_state(config, state)
                logging.info("Formatting complete: %s", pdf_path.name)

            final_notes = notes_path.read_text(encoding="utf-8")
            docx_notes_path = create_formatted_notes_docx(config, notes_path)
        else:
            logging.info("CREATE_STRUCTURED_NOTES=false, skipped fixed-template notes for %s.", pdf_path.name)

        if config.test_max_pages:
            logging.info("Test mode active, not appending partial notes to master and not marking PDF complete.")
            print(f"Verified source created: {source_path}")
            if docx_notes_path:
                print(f"Test Word document created: {docx_notes_path}")
            if verified_docx_path:
                print(f"Exact verified transcription Word document created: {verified_docx_path}")
            print("Partial test notes were not appended to Complete_Book_Notes.md.")
        elif has_failed_pages:
            logging.warning(
                "Not appending %s to master notes or marking it complete because failed pages still need review.",
                pdf_path.name,
            )
            print("DOCX created with manual-review placeholders. Run Again / Resume can retry the failed pages.")
        else:
            append_to_master_notes(config, state, pdf_path.name, final_notes)
            mark_pdf_complete(config, state, pdf_stem)
    except KeyboardInterrupt:
        raise
    except Exception as exc:
        logging.error("Chapter finalization failed for %s:\n%s", pdf_path.name, traceback.format_exc())
        state.setdefault("failed_pages", {}).setdefault(pdf_stem, {})["chapter_finalization"] = {
            "stage": "chapter_finalization",
            "error": str(exc),
            "timestamp": utc_now(),
        }
        save_state(config, state)
        print(f"Chapter finalization failed for {pdf_path.name}. Progress was saved.")


def print_summary(config: Config, state: dict[str, Any]) -> None:
    completed_pdfs = len(state.get("completed_pdfs", []))
    failed_pages = sum(
        len(pages) for pages in state.get("failed_pages", {}).values() if isinstance(pages, dict)
    )
    logging.info("Final summary: completed_pdfs=%s failed_entries=%s", completed_pdfs, failed_pages)
    print("\nSummary")
    print(f"Completed PDFs: {completed_pdfs}")
    print(f"Failed page/finalization entries: {failed_pages}")
    print(f"Master notes: {config.master_notes_file}")
    print(f"State file: {config.state_file}")
    print(f"Log file: {config.logs_dir / 'processor.log'}")


def main() -> None:
    config = load_config()
    setup_logging(config)
    ensure_output_dirs(config)
    print("Pipeline: Gemini multimodal transcription + verification for scanned textbook pages.")
    logging.info("Pipeline started: Gemini multimodal transcription + verification.")

    if not config.gemini_api_key:
        logging.error("Missing GEMINI_API_KEY. Add it to .env and run again.")
        print("Missing GEMINI_API_KEY. Create/edit .env and set GEMINI_API_KEY=your_key_here.")
        return

    state = load_state(config)
    if config.reset_master and not config.test_max_pages:
        config.master_notes_file.write_text("", encoding="utf-8")
        state["master_appended_chapters"] = []
        logging.info("RESET_MASTER=true, master notes file was cleared.")
    elif config.reset_master and config.test_max_pages:
        logging.info("RESET_MASTER=true ignored during test mode to protect master notes.")
        print("RESET_MASTER=true is ignored during test mode to protect Complete_Book_Notes.md.")
    save_state(config, state)
    pdfs = find_pdfs(config)
    if not pdfs:
        print(f"No PDF files found in {config.pdf_dir}. Add scanned PDFs and run again.")
        logging.warning("No PDFs found in %s", config.pdf_dir)
        return

    client = gemini_client(config)

    try:
        if config.test_pdf_limit:
            pdfs = pdfs[: config.test_pdf_limit]
            logging.info("TEST_PDF_LIMIT active: processing first %s PDF(s).", config.test_pdf_limit)
            print(f"Test PDF limit: processing first {len(pdfs)} PDF(s).")
        for pdf_path in pdfs:
            process_pdf(client, config, state, pdf_path)
    except KeyboardInterrupt:
        save_state(config, state)
        logging.warning("Interrupted by user. Progress saved.")
        print("\nInterrupted. Progress was saved in processing_state.json.")
    except FatalGeminiError as exc:
        save_state(config, state)
        logging.error("%s", exc)
        print(str(exc))
        raise SystemExit(2)
    except Exception:
        save_state(config, state)
        logging.error("Unexpected top-level error:\n%s", traceback.format_exc())
        print("Unexpected error. Progress was saved; check logs/processor.log.")
    finally:
        print_summary(config, state)


if __name__ == "__main__":
    main()
